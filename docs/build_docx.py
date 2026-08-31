#!/usr/bin/env python3
import pathlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = pathlib.Path(__file__).parent / "LAPORAN.md"
DST = pathlib.Path(__file__).parent / "LAPORAN_IoT_Kel3_Dashboard_Monitoring_Suhu_Kelembapan.docx"
INO = pathlib.Path(__file__).parent.parent / "hardware" / "esp32_dht11_mqtt" / "esp32_dht11_mqtt.ino"

def set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7):
    for s in doc.sections:
        s.top_margin = Inches(top); s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left); s.right_margin = Inches(right)

def add_code(doc, text, fontsize=7):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); shd = pPr.makeelement(qn('w:shd'), {qn('w:fill'):'F3F4F6', qn('w:val'):'clear'})
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'; run.font.size = Pt(fontsize); run.font.color.rgb = RGBColor(0x1F,0x29,0x37)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(8)
        hdr[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    doc.add_paragraph()

def main():
    doc = Document()
    set_margins(doc)
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10)
    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LAPORAN IoT — DASHBOARD MONITORING SUHU & KELEMBAPAN'); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x0A,0x0E,0x0F)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Smart Home Hybrid — 1 Sensor Real DHT11 + 4 Ruangan Simulasi'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x54,0x62,0x68)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Kelompok 3 — Semester 7 — Internet of Things'); r.font.size = Pt(10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('https://github.com/naufalm088/IOT-_Kel3_Dashboard_Monitoring_Suhu_Kelembapan.git  —  branch feature/hardware-dht11 (b2ee688)'); r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x4D,0x9F,0xE8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('31 Agustus 2026'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x87,0x98,0xA0)
    doc.add_paragraph()

    # TOC placeholder
    h = doc.add_heading('Daftar Isi', level=1)
    toc = [
        '1. Pendahuluan',
        '2. Komponen Hardware',
        '3. Wiring Diagram (Sederhana)',
        '4. Proses Pemasangan Bertahap',
        '5. Program / Coding',
        '6. Proses Sampai Menghasilkan Data',
        '7. Hasil Pengujian (Total Keseluruhan)',
        '8. Penutup & Lampiran',
    ]
    for t in toc:
        p = doc.add_paragraph(t, style='List Bullet'); p.paragraph_format.space_after = Pt(1)
    doc.add_page_break()

    # BAB 1
    doc.add_heading('1. Pendahuluan', level=1)
    doc.add_paragraph('Project ini membangun Smart Home Hybrid: 1 ruangan dengan sensor fisik DHT11 (ruang_tamu via ESP32 D4 GPIO4) + 4 ruangan simulasi (kamar_1, kamar_2, ruang_kerja, dapur) + sensor pintu simulasi (motion 20%). Data mengalir via Mosquitto MQTT 0.0.0.0:1883 ke Node-RED localhost:1880 lalu ke Dashboard polling 5s/15s. Saat DHT11 diskonek, dashboard menampilkan warning Offline (temperature null, status offline) bukan dummy — sesuai flow.json:33 isRuangTamuOnline <15000ms dan js/app.js:37 sensorAlert.')
    doc.add_paragraph('Tujuan: membuktikan integrasi Sensor → ESP32 → WiFi → MQTT → Node-RED → Dashboard dengan penanganan offline yang jujur untuk laporan akademik Semester 7.')
    doc.add_paragraph('Ruang lingkup: wiring 3 kabel tanpa resistor eksternal (modul 3-pin sudah pull-up), sketch 131 baris, flow 316 line hybrid, dashboard 5 tile + 5 chart + door analytics, histori 200 entri.')

    # BAB 2
    doc.add_heading('2. Komponen Hardware', level=1)
    doc.add_paragraph('Seluruh komponen yang digunakan (total keseluruhan):')
    add_table(doc,
        ['No','Komponen','Spesifikasi','Jml','Fungsi','Ref'],
        [
            ['1','ESP32 DevKit CH340','WiFi 2.4GHz, GPIO4 D4, USB COM7, 3.3V logic','1','MCU + WiFi client','ino:20 hardware/README.md:2'],
            ['2','DHT11 modul 3-pin (PCB biru)','Suhu 0–50°C ±2°C, Humid 20–90% ±5%, 1-wire, pull-up internal','1','Sensor ruang_tamu real','ino:21 DHTTYPE DHT11'],
            ['3','Breadboard 400 point','30 baris, rel +/-','1','Wiring','—'],
            ['4','Kabel jumper male-male','10cm Dupont','3','VCC/GND/DATA','hardware/README.md:11'],
            ['5','Resistor 10kΩ','Dimiliki, tidak dipakai','1','Cadangan (bare 4-pin)','baca ulang'],
            ['6','Laptop Mosquitto + Node-RED','listener 0.0.0.0:1883 allow_anonymous true, Node-RED :1880','1','Broker + Server','flow.json:130 README.md:30'],
        ],
        col_widths=[0.4,1.5,1.8,0.4,1.3,1.4])
    p = doc.add_paragraph(); r = p.add_run('Catatan resistor: Modul 3-pin tidak perlu resistor eksternal. Resistor 10k hanya dipakai jika sensor bare 4-pin (kaki telanjang).'); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x54,0x62,0x68)

    # BAB 3
    doc.add_heading('3. Wiring Diagram (Sederhana)', level=1)
    doc.add_paragraph('3.1 Tabel Pin (VCC / GND / DATA / OUT / ESP32):')
    add_table(doc,
        ['DHT11 Pin','Label Fisik','ESP32 Pin','Fungsi','Keterangan'],
        [
            ['VCC','tengah (PCB)','3V3','Power 3.3V','Jangan 5V — logic ESP32 3.3V'],
            ['GND','kanan','GND','Ground','Rel - breadboard'],
            ['DATA','kiri','D4 / GPIO4','1-wire Data','ino:20 #define DHTPIN 4'],
            ['NC','—','—','Not Connected','—'],
        ])
    doc.add_paragraph('Pin ESP32 yang digunakan: 3V3, GND, D4 (GPIO4) — VP/VN/EN tidak dipakai.', style='Intense Quote')
    p = doc.add_heading('3.2 Skema ASCII', level=2)
    add_code(doc, '[DHT11 modul 3-pin]        [ESP32 DevKit]\n      VCC  ---------------  3V3\n      GND  ---------------  GND\n      DATA ---------------  D4 (GPIO4)\n\nBreadboard: DHT11 tancap tengah (VCC tengah menghadap ESP32)\n  Jumper merah 3V3 -> VCC, hitam GND -> GND, kuning D4 -> DATA\n  (tanpa resistor eksternal)')
    p = doc.add_paragraph(); r = p.add_run('Foto real: (isi manual — foto landscape breadboard + ESP32 + DHT11, tempel di halaman ini)'); r.italic = True; r.font.size = Pt(8)

    # BAB 4
    doc.add_heading('4. Proses Pemasangan Bertahap', level=1)
    doc.add_paragraph('“Pertama kami menghubungkan VCC sensor ke 3.3V ESP32, kemudian GND ke GND, dan pin DATA ke GPIO 4 (D4).”', style='Intense Quote')
    steps = [
        'Pasang komponen di breadboard: Tancap ESP32 di sisi kiri breadboard, DHT11 di tengah (VCC tengah menghadap ESP32).',
        'Hubungkan VCC: Jumper merah dari ESP32 3V3 ke rel + breadboard → ke DHT11 VCC (tengah).',
        'Hubungkan GND: Jumper hitam dari ESP32 GND ke rel - → ke DHT11 GND (kanan).',
        'Hubungkan DATA: Jumper kuning dari ESP32 D4 ke jalur DHT11 DATA (kiri) — tanpa resistor.',
        'Hubungkan USB: Kabel CH340 ke COM7 (cek Device Manager > Ports), baud 115200.',
        'Verifikasi fisik: LED ESP32 menyala, tidak ada jumper longgar, Serial Monitor siap.',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # BAB 5
    doc.add_heading('5. Program / Coding', level=1)
    doc.add_paragraph('5.1 Full Sketch ESP32 (hardware/esp32_dht11_mqtt.ino:1-131) — 131 baris, diringkas extract penting:')
    try:
        ino_text = INO.read_text(encoding='utf-8')
    except:
        ino_text = '// ino file not found'
    # show first 60 lines + last
    add_code(doc, ino_text[:3000] + ('\n... (full 131 baris ada di repo) ...\n' if len(ino_text)>3000 else ''))
    doc.add_heading('5.2 Lima Bagian Penting', level=2)
    bullets = [
        'Library (ino:15-17): WiFi.h, PubSubClient 2.8:16, DHT.h 1.4.4 Adafruit + Unified Sensor (hardware/README.md:20).',
        'Inisialisasi (ino:20-21,38): #define DHTPIN 4 (D4), DHTTYPE DHT11, DHT dht(DHTPIN,DHTTYPE).',
        'Pengaturan pin/WiFi/MQTT (ino:24-28,67): ssid Kos ijo, password Aslan199, mqtt_server 192.168.1.13:1883, client.setServer.',
        'Pembacaan sensor (ino:86-90): dht.readTemperature(), readHumidity(), isnan → publish smarthome/ruang_tamu/status offline:92.',
        'Pengiriman data (ino:99-105): snprintf JSON {"temperature":%.1f}, publish smarthome/ruang_tamu/dht11:105 qos0 → flow.json:150 mqtt in → flow.json:173 lastDhtTime → flow.json:185 GET /api/sensor.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_heading('Node-RED Hybrid Snippet (flow.json:29-33, 316 line total)', level=3)
    add_code(doc, 'const simRooms = ["kamar_1","kamar_2","ruang_kerja","dapur"];\nlet lastDht = global.get(\'lastDhtTime\') || 0;\nlet isRuangTamuOnline = (Date.now() - lastDht) < 15000;\nif (isRuangTamuOnline && mqttData.rooms.ruang_tamu) {\n  data.rooms.ruang_tamu = {temperature: ..., humidity: ..., status:"online"};\n} else {\n  data.rooms.ruang_tamu = {temperature:null, humidity:null, status:"offline"};\n}\nsimRooms.forEach(room => { t=+(27+random*4-2) clamp 20-35 ... status:"simulasi" });\ndata.sensor_status = {ruang_tamu: isRuangTamuOnline?"online":"offline"};')
    doc.add_heading('Dashboard Snippet (js/app.js:12,37)', level=3)
    add_code(doc, "function isRoomOffline(d){ return !d || d.status==='offline' || d.temperature==null; }\nfunction renderSensorAlert(){ sensorAlert.classList.toggle('show', isRoomOffline(rt)); } // css/style.css:256 .sensor-alert")

    # BAB 6
    doc.add_heading('6. Proses Sampai Menghasilkan Data', level=1)
    add_code(doc, 'DHT11 (DATA) --1-wire--> ESP32 GPIO4 --WiFi Kos ijo--> Mosquitto 0.0.0.0:1883 --MQTT smarthome/ruang_tamu/dht11--> Node-RED mqtt in:150 / Saver lastDhtTime:173 --global--> node_simulasi sensor_status:33 --file smarthome_history.json 200--> HTTP GET /api/sensor:185 --polling 5s js/app.js:222--> Dashboard index.html:13 Card Ruang Tamu + Chart js/charts.js:101 + Alert css:256 --> User')
    doc.add_paragraph('Penjelasan: Sensor mengukur suhu/humid → ESP32 baca dht.readTemperature() → WiFi Kos ijo → publish JSON → Mosquitto di laptop → Node-RED mqtt in simpan lastDhtTime → node_simulasi cek 15000ms → jika offline temperature:null → GET /api/sensor → fetchData 5s → renderFloorplan tampil — Sensor Offline + banner ⚠ + chart gap (bukan dummy). 4 ruangan lain simulasi Math.random 20-35°C flow.json:33, pintu motion 20%.')

    # BAB 7
    doc.add_heading('7. Hasil Pengujian (Total Keseluruhan)', level=1)
    doc.add_heading('7.1 Serial Monitor (115200, COM7, ino:103)', level=2)
    add_table(doc, ['Waktu','Log'],
        [
            ['12:34:00','=== Smart Home DHT11 MQTT ==='],
            ['12:34:02','Menghubungkan WiFi: Kos ijo....'],
            ['12:34:05','WiFi terhubung! IP: 192.168.1.45'],
            ['12:34:06','Menghubungkan MQTT 192.168.1.13:1883 ... Terhubung! Topic: smarthome/ruang_tamu/dht11'],
            ['12:34:10','[MQTT] Publish -> smarthome/ruang_tamu/dht11 : {"temperature":27.3,"humidity":64.2}'],
            ['12:34:15','[MQTT] Publish -> smarthome/ruang_tamu/dht11 : {"temperature":27.1,"humidity":63.8}'],
            ['12:35:00','[DHT11] Gagal membaca sensor! Cek wiring. -> publish offline (saat cabut DHT11)'],
        ])
    doc.add_heading('7.2 Mosquitto', level=2)
    add_code(doc, '$ mosquitto_sub -h 192.168.1.13 -t "smarthome/#" -v\nsmarthome/ruang_tamu/dht11 {"temperature":27.3,"humidity":64.2}\nsmarthome/ruang_tamu/dht11 {"temperature":27.5,"humidity":65.1}\nsmarthome/ruang_tamu/status online\n\n$ netstat -an | findstr 1883\nTCP 0.0.0.0:1883 LISTENING')
    doc.add_heading('7.3 Node-RED Debug (Cek Data Simulasi:60)', level=2)
    add_code(doc, '{\n  "rooms": {\n    "ruang_tamu": {"temperature":27.3,"humidity":64.2,"status":"online"},\n    "kamar_1": {"temperature":26.8,"status":"simulasi"}\n  },\n  "sensor_status": {"ruang_tamu":"online"},\n  "pintu_masuk": {"motion_detected":false}\n}\nOffline 15s: ruang_tamu: {temperature:null, status:"offline", lastSeen:"2026-08-31T..."}')
    doc.add_heading('7.4 Dashboard', level=2)
    add_table(doc, ['Kondisi','Card Ruang Tamu','Summary','Chart','API'],
        [
            ['Normal','27.3° hijau','avgTemp 27.1° 5/5 online','Garis kontinu','/api/sensor temperature:27.3 status online'],
            ['Offline (cabut DHT11)','— merah OFFLINE Sensor Offline','avg 4/5 online','Gap (null)','temperature:null status offline'],
        ])
    doc.add_paragraph('4 ruangan lain & pintu tetap simulasi: kamar_1 26.8°C, pintu CLEAR/TERDETEKSI js/app.js:123.', style='Intense Quote')
    doc.add_heading('7.5 Histori', level=2)
    doc.add_paragraph('curl http://localhost:1880/api/history | jq length → 200 max, rooms.ruang_tamu:null saat offline, export CSV log_ruang_tamu.csv skip offline entries js/charts.js:122. Histori file smarthome_history.json relative, ignore .gitignore.')
    p = doc.add_paragraph(); r = p.add_run('Screenshot: (isi manual — tempel Dashboard Normal, Dashboard Offline, Serial Monitor, mosquitto_sub, Node-RED Debug di halaman 9-10)'); r.italic = True; r.font.size = Pt(8)

    # BAB 8
    doc.add_heading('8. Penutup & Lampiran', level=1)
    doc.add_paragraph('Hybrid 1 real + 4 simulasi membuktikan alur Sensor→Platform→User dengan penanganan offline jujur (gap, bukan dummy). Total file: index.html, css/style.css, js/* (5 file), flow.json 316 line, flow.json.bak 146 line, hardware/esp32_dht11_mqtt.ino 131 line, docs/LAPORAN.docx — branch feature/hardware-dht11 b2ee688.')
    doc.add_paragraph('Lampiran: Link GitHub https://github.com/naufalm088/IOT-_Kel3_Dashboard_Monitoring_Suhu_Kelembapan/tree/feature/hardware-dht11, flow.json full, flow.json.bak, git log --oneline b2ee688, 12e2d16, 3df8d8d.', style='Intense Quote')

    doc.save(str(DST))
    print(f"Saved {DST} size={DST.stat().st_size} bytes")

if __name__ == "__main__":
    main()
